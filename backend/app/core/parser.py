import pdfplumber
import fitz  # PyMuPDF fallback
import io
import re
from docx import Document
from fastapi import HTTPException, status
from app.core.nlp_engine import nlp

# ── Section Mapping ──────────────────────────────────────────────────────────

SECTION_MAP = {
    "summary": ["summary", "objective", "professional summary", "about me", "profile", "professional profile"],
    "experience": ["experience", "employment", "work history", "professional experience", "work experience", "career history", "internships", "internship"],
    "skills": ["skills", "technical skills", "core competencies", "technologies", "key skills", "expertise"],
    "projects": ["projects", "personal projects", "academic projects", "key projects", "technical projects"],
    "education": ["education", "academic background", "qualifications", "academic profile", "academic credentials"],
    "contact": ["contact", "personal info", "contact information"]
}

# ── Header Detection ─────────────────────────────────────────────────────────

def is_header(line_text: str, font_size: float = 0, is_bold: bool = False) -> bool:
    """
    Heuristic to determine if a line is a section header.
    - No verbs (POS tagging)
    - Short length (usually < 5 words)
    - High capitalization or bold/large font
    """
    text = line_text.strip()
    if not text or len(text) > 50:
        return False
    
    # Short-circuit for extremely common keywords (always headers if short)
    if text.lower() in ["education", "skills", "experience", "projects", "summary", "contact"]:
        return True
        
    # Pass 1: POS Tagging - Lines with no VERB are likely headers
    doc = nlp(text)
    has_verb = any(token.pos_ == "VERB" for token in doc)
    if has_verb:
        return False

    # Pass 1: Visual/Textual heuristics
    word_count = len(text.split())
    if word_count > 6:
        return False

    if text.isupper() and word_count < 4:
        return True
        
    if (is_bold or font_size > 11.5) and word_count < 5:
        return True
        
    return False

# ── PDF Extraction (Two-Pass) ────────────────────────────────────────────────

def extract_sections_from_pdf(file_bytes: bytes) -> dict:
    """
    Extracts text using pdfplumber as primary with a two-pass approach:
    1. Line classification (Header vs Content)
    2. Bucketing into semantic sections
    """
    sections = {k: [] for k in SECTION_MAP.keys()}
    sections["other"] = []
    current_section = "other"
    
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Extract words with visual attributes
                words = page.extract_words(extra_attrs=["fontname", "size"])
                if not words:
                    continue

                # Group words into lines based on vertical position (top)
                lines = {}
                for w in words:
                    top = round(w["top"], 1)
                    if top not in lines:
                        lines[top] = []
                    lines[top].append(w)
                
                # Process lines in order
                last_bottom = 0
                for top in sorted(lines.keys()):
                    line_words = sorted(lines[top], key=lambda x: x["x0"])
                    line_text = " ".join(w["text"] for w in line_words).strip()
                    
                    if not line_text:
                        continue
                        
                    avg_size = sum(w["size"] for w in line_words) / len(line_words)
                    is_bold = any(re.search(r"bold|heavy|black", w["fontname"].lower()) for w in line_words)
                    
                    # Pass 2: Header detection and section switching
                    if is_header(line_text, avg_size, is_bold):
                        clean_header = line_text.lower().strip()
                        found_new_section = False
                        for sec, keywords in SECTION_MAP.items():
                            if any(k == clean_header or (len(clean_header) < 20 and k in clean_header) for k in keywords):
                                current_section = sec
                                found_new_section = True
                                break
                        
                        # If it's a header but not in our map, keep it in "other" or current
                        if not found_new_section:
                            sections[current_section].append(line_text)
                    else:
                        sections[current_section].append(line_text)

        # Post-process: Join lines and remove empty sections
        result = {k: "\n".join(v).strip() for k, v in sections.items() if v}
        
        # If result is too empty, something went wrong (likely scanned PDF)
        if len("".join(result.values())) < 100:
            return _fallback_py_mu_pdf(file_bytes)
            
        return result

    except Exception:
        return _fallback_py_mu_pdf(file_bytes)

def _fallback_py_mu_pdf(file_bytes: bytes) -> dict:
    """Fallback using PyMuPDF (fitz) for raw text if pdfplumber fails."""
    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in pdf)
        pdf.close()
        return {"other": text, "is_fallback": True}
    except Exception:
        return {"other": ""}

# ── DOCX Extraction ──────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> dict:
    """Extracts text from DOCX. Currently simple, but returns dict for consistency."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        
        return {"other": "\n".join(parts)}
    except Exception:
        return {"other": ""}

# ── Unified Registry ──────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, file_ext: str) -> dict:
    """
    Main entry point for file parsing.
    Returns a dictionary of sections: { "skills": "...", "experience": "...", ... }
    """
    if file_ext == "pdf":
        return extract_sections_from_pdf(file_bytes)
    elif file_ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {file_ext}"
        )